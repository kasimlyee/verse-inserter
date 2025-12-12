"""
Pytest configuration and shared fixtures for VerseInserter tests.

Provides common fixtures, test data, and configuration for all test modules.
"""

import pytest
from pathlib import Path
from typing import Generator
from docx import Document
from unittest.mock import Mock, AsyncMock

from verse_inserter.models.verse import Verse, VerseReference, TranslationType
from verse_inserter.config.settings import Settings


@pytest.fixture
def sample_verse_reference() -> VerseReference:
    """Create a sample verse reference for testing."""
    return VerseReference(
        book="John",
        chapter=3,
        start_verse=16,
        translation=TranslationType.NIV
    )


@pytest.fixture
def sample_verse_range() -> VerseReference:
    """Create a sample verse range for testing."""
    return VerseReference(
        book="Psalm",
        chapter=23,
        start_verse=1,
        end_verse=3,
        translation=TranslationType.KJV
    )


@pytest.fixture
def sample_verse(sample_verse_reference: VerseReference) -> Verse:
    """Create a sample verse with text for testing."""
    return Verse(
        reference=sample_verse_reference,
        text="For God so loved the world that he gave his one and only Son, that whoever believes in him shall not perish but have eternal life.",
        translation=TranslationType.NIV,
        source_api="test_api"
    )


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Create a sample Word document for testing."""
    doc_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It contains {{John 3:16}} and {{Psalm 23:1-3}}.")
    doc.add_paragraph("Another verse: {{Genesis 1:1}}.")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def sample_docx_with_tables(tmp_path: Path) -> Path:
    """Create a Word document with tables containing placeholders."""
    doc_path = tmp_path / "test_document_tables.docx"
    doc = Document()
    doc.add_paragraph("Document with table:")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Reference"
    table.cell(0, 1).text = "{{John 3:16}}"
    table.cell(1, 0).text = "Another"
    table.cell(1, 1).text = "{{Psalm 23:1}}"

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def mock_api_client() -> AsyncMock:
    """Create a mock API client for testing."""
    mock_client = AsyncMock()

    async def mock_fetch_verse(reference: VerseReference) -> Verse:
        return Verse(
            reference=reference,
            text="Mocked verse text for testing",
            translation=reference.translation,
            source_api="mock_api"
        )

    mock_client.fetch_verse.side_effect = mock_fetch_verse
    mock_client.fetch_verse_with_fallback.side_effect = mock_fetch_verse

    return mock_client


@pytest.fixture
def mock_settings(tmp_path: Path) -> Settings:
    """Create mock settings for testing."""
    settings = Settings()
    settings.api_key = "test_api_key_12345"
    settings.nlt_api_key = "test_nlt_key_67890"
    settings.default_translation = "NIV"
    settings.enable_cache = False  # Disable cache for testing
    return settings


@pytest.fixture
def clean_cache(tmp_path: Path) -> Generator[Path, None, None]:
    """Provide a clean cache directory for testing."""
    cache_dir = tmp_path / "test_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    yield cache_dir
    # Cleanup after test
    import shutil
    if cache_dir.exists():
        shutil.rmtree(cache_dir)


# Test data constants
SAMPLE_PLACEHOLDERS = [
    "{{John 3:16}}",
    "{{Genesis 1:1}}",
    "{{Psalm 23:1-6}}",
    "{{1 Corinthians 13:4-7}}",
    "{{Revelation 21:4}}",
]

SAMPLE_BOOK_NAMES = [
    ("Genesis", "GEN"),
    ("John", "JHN"),
    ("1 Corinthians", "1CO"),
    ("Psalm", "PSA"),
    ("Revelation", "REV"),
]

INVALID_REFERENCES = [
    "John 3",  # Missing verse
    "Invalid Book 1:1",  # Invalid book name
    "John 3:999",  # Verse out of range
    "{{Unclosed placeholder",  # Malformed
    "John 3:16-10",  # Invalid range (end < start)
]
