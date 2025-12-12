"""
Tests for formatting preservation during text replacement.

Verifies that run-level formatting (bold, italic, font, color) is preserved
when replacing placeholders with verse text.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

from verse_inserter.core.document_processor import DocumentProcessor
from verse_inserter.models.verse import Verse, VerseReference, TranslationType


class TestFormattingPreservation:
    """Tests for formatting preservation during replacement."""

    @pytest.fixture
    def docx_with_formatting(self, tmp_path: Path) -> Path:
        """Create a document with formatted text containing placeholders."""
        doc_path = tmp_path / "formatted_doc.docx"
        doc = Document()

        # Add paragraph with mixed formatting
        para = doc.add_paragraph()

        # Add normal text
        run1 = para.add_run("This is ")

        # Add bold placeholder
        run2 = para.add_run("{{John 3:16}}")
        run2.bold = True
        run2.font.size = Pt(14)

        # Add normal text after
        run3 = para.add_run(" and more text.")

        # Add another paragraph with italic placeholder
        para2 = doc.add_paragraph()
        run4 = para2.add_run("Reference: ")
        run5 = para2.add_run("{{Genesis 1:1}}")
        run5.italic = True
        run5.font.color.rgb = RGBColor(255, 0, 0)  # Red

        doc.save(str(doc_path))
        return doc_path

    def test_preserve_bold_formatting(self, docx_with_formatting):
        """Test that bold formatting is preserved during replacement."""
        processor = DocumentProcessor(create_backup=False, preserve_formatting=True)

        verse = Verse(
            reference=VerseReference.parse("John 3:16"),
            text="For God so loved the world...",
            translation=TranslationType.NIV
        )

        with processor.load_document(docx_with_formatting) as doc:
            # Get original paragraph
            para = doc.paragraphs[0]

            # Find the bold run
            bold_run = None
            for run in para.runs:
                if run.bold and "John 3:16" in run.text:
                    bold_run = run
                    break

            assert bold_run is not None, "Bold placeholder run not found"
            original_font_size = bold_run.font.size

            # Perform replacement
            verse_map = {"John 3:16": verse}
            processor.replace_placeholders(doc, verse_map)

            # Check that formatting is preserved
            # The run containing the verse should still be bold
            verse_run = None
            for run in para.runs:
                if "For God so loved" in run.text:
                    verse_run = run
                    break

            # Verify run exists and has formatting preserved
            assert verse_run is not None, "Verse run not found after replacement"
            assert verse_run.bold is True, "Bold formatting was lost"
            assert verse_run.font.size == original_font_size, "Font size was changed"

    def test_preserve_italic_formatting(self, docx_with_formatting):
        """Test that italic formatting is preserved during replacement."""
        processor = DocumentProcessor(create_backup=False, preserve_formatting=True)

        verse = Verse(
            reference=VerseReference.parse("Genesis 1:1"),
            text="In the beginning God created...",
            translation=TranslationType.NIV
        )

        with processor.load_document(docx_with_formatting) as doc:
            # Get original paragraph
            para = doc.paragraphs[1]

            # Find the italic run
            italic_run = None
            for run in para.runs:
                if run.italic and "Genesis 1:1" in run.text:
                    italic_run = run
                    break

            assert italic_run is not None, "Italic placeholder run not found"

            # Perform replacement
            verse_map = {"Genesis 1:1": verse}
            processor.replace_placeholders(doc, verse_map)

            # Check that formatting is preserved
            verse_run = None
            for run in para.runs:
                if "In the beginning" in run.text:
                    verse_run = run
                    break

            assert verse_run is not None, "Verse run not found after replacement"
            assert verse_run.italic is True, "Italic formatting was lost"

    def test_preserve_surrounding_text(self, docx_with_formatting):
        """Test that text before and after placeholder is preserved."""
        processor = DocumentProcessor(create_backup=False, preserve_formatting=True)

        verse = Verse(
            reference=VerseReference.parse("John 3:16"),
            text="For God so loved the world...",
            translation=TranslationType.NIV
        )

        with processor.load_document(docx_with_formatting) as doc:
            verse_map = {"John 3:16": verse}
            processor.replace_placeholders(doc, verse_map)

            # Check full text is correct
            para_text = doc.paragraphs[0].text
            assert "This is " in para_text
            assert "For God so loved the world..." in para_text
            assert "and more text." in para_text

    def test_replacement_across_multiple_runs(self, tmp_path):
        """Test replacement when placeholder spans multiple runs."""
        doc_path = tmp_path / "multi_run_doc.docx"
        doc = Document()

        para = doc.add_paragraph()
        para.add_run("Start ")
        para.add_run("{{John")  # Placeholder split across runs
        para.add_run(" 3:16}}")
        para.add_run(" End")

        doc.save(str(doc_path))

        processor = DocumentProcessor(create_backup=False, preserve_formatting=True)
        verse = Verse(
            reference=VerseReference.parse("John 3:16"),
            text="Verse text",
            translation=TranslationType.NIV
        )

        with processor.load_document(doc_path) as doc:
            verse_map = {"John 3:16": verse}
            result = processor.replace_placeholders(doc, verse_map)

            # Verify replacement worked
            para_text = doc.paragraphs[0].text
            assert "Verse text" in para_text
            assert "{{John 3:16}}" not in para_text

    def test_legacy_mode_without_preservation(self, docx_with_formatting):
        """Test that legacy mode still works (preserve_formatting=False)."""
        processor = DocumentProcessor(create_backup=False, preserve_formatting=False)

        verse = Verse(
            reference=VerseReference.parse("John 3:16"),
            text="For God so loved the world...",
            translation=TranslationType.NIV
        )

        with processor.load_document(docx_with_formatting) as doc:
            verse_map = {"John 3:16": verse}
            result = processor.replace_placeholders(doc, verse_map)

            # Verify replacement happened
            assert result.placeholders_replaced > 0
            para_text = doc.paragraphs[0].text
            assert "For God so loved the world..." in para_text
